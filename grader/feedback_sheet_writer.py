"""Helpers for populating feedback-sheet DOCX templates."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from grader.agents.report_agent import q2_page_range_text
from grader.models import FinalGrade, ScoreValue
from grader.utils import compact_text, format_score

ROW_LABEL_1A = compact_text("All worksheets included in OscorpCleaned.xlsx")
ROW_LABEL_1B = compact_text("Final data set properly cleaned")
ROW_LABEL_1C = compact_text("Pivot table for report 1")
ROW_LABEL_1D = compact_text("Pivot table for report 2")
ROW_LABEL_Q1 = compact_text("Response to reflection question #1")
ROW_LABEL_Q2 = compact_text("Response to reflection question #2")
ROW_LABEL_TOTAL = compact_text("TOTAL")

STEP1B_GUIDANCE_DEFAULT = "Review this section and make the needed updates in your cleaned dataset."


def to_student_facing_step1b_notes(errors_1b: list[str]) -> list[str]:
    if not errors_1b:
        return ["No updates needed for this section."]

    guidance: list[str] = []
    for error in errors_1b:
        text = error.strip()
        if not text:
            continue

        if text == "No errors found.":
            guidance.append("No updates needed for this section.")
            continue

        if text.startswith("Required 17-column structure"):
            missing_label = "Missing:"
            if missing_label in text:
                missing_columns = text.split(missing_label, 1)[1].strip().rstrip(".")
                guidance.append(f"Include these columns in your cleaned dataset: {missing_columns}.")
            else:
                guidance.append("Include all required Step 1B columns in your cleaned dataset.")
            continue

        if text.startswith("Duplicate transaction rows remain"):
            guidance.append("Remove duplicate transaction rows so each transaction appears only once.")
            continue

        if text.startswith("TransactionType still contains"):
            guidance.append(
                "Use full-word TransactionType values only: Fee, Deposit, Withdrawal, Transfer, Payment, Transaction."
            )
            continue

        if text.startswith("TransactionID values are not consistently formatted"):
            guidance.append("Format each TransactionID as T#### (example: T0001).")
            continue

        if text.startswith("CustomerFName and CustomerLName are not properly split"):
            guidance.append("Split customer names into CustomerFName and CustomerLName columns.")
            continue

        if text.startswith("OpenDate column is missing"):
            guidance.append("Add an OpenDate column and include date and time values.")
            continue

        if text.startswith("OpenDate is missing a time component"):
            guidance.append("Include both date and time in OpenDate for each filled row.")
            continue

        if text.startswith("Merchant still contains city/state text"):
            guidance.append("Keep merchant names in Merchant and place location values in City and State.")
            continue

        if text.startswith("City/State values appear unsplit"):
            guidance.append("Enter City and State in their own columns when location information is provided.")
            continue

        if text.startswith("TotalValue is not consistently stored"):
            guidance.append("Store TotalValue as numeric currency values, not text.")
            continue

        if text.startswith("Notes still contains special characters"):
            guidance.append("Update Notes to remove unresolved err-style text and unnecessary special characters.")
            continue

        if text.startswith("Explicit null-style markers remain"):
            guidance.append("Replace null-style markers with cleaned values, or leave the cell blank when appropriate.")
            continue

        guidance.append(STEP1B_GUIDANCE_DEFAULT)

    if not guidance:
        return ["No updates needed for this section."]

    return guidance


def feedback_template_key(filename: str) -> str:
    stem = Path(filename).stem
    prefix = "feedback_sheet_"
    lower = stem.lower()
    if lower.startswith(prefix):
        stem = stem[len(prefix) :]
    return compact_text(stem.replace("_", ""))


def build_feedback_template_index(feedback_sheets_dir: Path) -> dict[str, Path]:
    index: dict[str, Path] = {}
    if not feedback_sheets_dir.exists():
        return index

    for path in sorted(feedback_sheets_dir.glob("*.docx")):
        key = feedback_template_key(path.name)
        if key and key not in index:
            index[key] = path

    return index


def write_feedback_sheet(template_path: Path, output_path: Path, grade: FinalGrade) -> None:
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError(f"Feedback sheet template has no table: {template_path}")

    updates = row_updates_for_grade(grade)
    table = doc.tables[0]

    for row in table.rows:
        if len(row.cells) < 4:
            continue
        row_label = compact_text(row.cells[0].text)
        if row_label not in updates:
            continue
        score_text, note_text = updates[row_label]
        row.cells[2].text = score_text
        row.cells[3].text = note_text

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def write_feedback_sheet_for_student(
    student_key: str,
    grade: FinalGrade,
    feedback_sheets_dir: Path,
    output_path: Path,
) -> Path | None:
    template_index = build_feedback_template_index(feedback_sheets_dir)
    template_path = template_index.get(student_key)
    if template_path is None:
        return None

    write_feedback_sheet(template_path, output_path, grade)
    return template_path


def write_feedback_sheets_for_batch(
    grades: dict[str, FinalGrade],
    feedback_sheets_dir: Path,
    output_dir: Path,
) -> tuple[int, list[str]]:
    template_index = build_feedback_template_index(feedback_sheets_dir)
    written = 0
    missing_templates: list[str] = []

    for key, grade in sorted(grades.items()):
        template_path = template_index.get(key)
        if template_path is None:
            missing_templates.append(key)
            continue

        output_path = output_dir / template_path.name
        write_feedback_sheet(template_path, output_path, grade)
        written += 1

    return written, missing_templates


def row_updates_for_grade(grade: FinalGrade) -> dict[str, tuple[str, str]]:
    x = grade.xlsx_eval
    r = grade.reflection_eval

    notes_1a = "" if is_full_credit(x.score_1a, 1.0) else x.notes_1a

    if is_full_credit(x.score_1b, 4.0):
        notes_1b = ""
    else:
        notes_1b_lines = to_student_facing_step1b_notes(x.errors_1b)
        notes_1b = "\n".join(f"- {item}" for item in notes_1b_lines)

    notes_1c = "" if is_full_credit(x.score_1c, 0.5) else x.notes_1c
    notes_1d = "" if is_full_credit(x.score_1d, 0.5) else x.notes_1d

    q1_notes = "" if is_full_credit(r.q1.score, 2.0) else f"{r.q1.label}: {r.q1.rationale}".strip()
    q2_pages = q2_page_range_text(r.q2_start_page, r.q2_end_page)
    q2_notes = "" if is_full_credit(r.q2.score, 2.0) else f"{r.q2.label}: {r.q2.rationale}\nSource pages: {q2_pages}".strip()

    total_notes = ""

    return {
        ROW_LABEL_1A: (score_text(x.score_1a), notes_1a),
        ROW_LABEL_1B: (score_text(x.score_1b), notes_1b),
        ROW_LABEL_1C: (score_text(x.score_1c), notes_1c),
        ROW_LABEL_1D: (score_text(x.score_1d), notes_1d),
        ROW_LABEL_Q1: (score_text(r.q1.score), q1_notes),
        ROW_LABEL_Q2: (score_text(r.q2.score), q2_notes),
        ROW_LABEL_TOTAL: (score_text(grade.total_score), total_notes),
    }


def score_text(score: ScoreValue) -> str:
    if isinstance(score, str):
        return score
    return format_score(float(score))


def is_full_credit(score: ScoreValue, max_score: float) -> bool:
    if isinstance(score, str):
        return False
    return abs(float(score) - max_score) < 1e-9