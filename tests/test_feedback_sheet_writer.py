"""Unit tests for feedback sheet DOCX population."""

from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from grader.feedback_sheet_writer import (
    feedback_template_key,
    to_student_facing_step1b_notes,
    write_feedback_sheet_for_student,
    write_feedback_sheets_for_batch,
)
from grader.models import (
    FinalGrade,
    ReflectionEvaluation,
    ReflectionQuestionResult,
    XlsxEvaluation,
)


ROW_LABELS = [
    "Item",
    "All worksheets included in OscorpCleaned.xlsx",
    "Final data set properly cleaned",
    "Pivot table for report 1",
    "Pivot table for report 2",
    "Response to reflection question #1",
    "Response to reflection question #2",
    "TOTAL",
]

POSSIBLE_SCORES = ["Possible Score", "1", "4", "0.5", "0.5", "2", "2", "10"]


def build_template(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=len(ROW_LABELS), cols=4)
    for row_index, label in enumerate(ROW_LABELS):
        table.rows[row_index].cells[0].text = label
        table.rows[row_index].cells[1].text = POSSIBLE_SCORES[row_index]
    doc.save(path)


def sample_grade() -> FinalGrade:
    xlsx_eval = XlsxEvaluation(
        score_1a=1.0,
        notes_1a="All seven worksheets present.",
        score_1b=3.5,
        errors_1b=["Duplicate transaction rows remain."],
        review_flags_1b=["Currency formatting confidence was low."],
        score_1c=0.5,
        notes_1c="Pivot object exists.",
        score_1d=0.5,
        notes_1d="Pivot object exists.",
        review_flags=[],
    )
    reflection_eval = ReflectionEvaluation(
        q1=ReflectionQuestionResult(
            score=1.25,
            label="Meets Expectations",
            rationale="Generally correct with limited depth.",
        ),
        q2=ReflectionQuestionResult(
            score=0.75,
            label="Does Not Meet Expectations",
            rationale="Too vague and missing concrete impact.",
        ),
        q2_start_page=2,
        q2_end_page=3,
        review_flags=[],
    )
    return FinalGrade(
        student_name="Bagley Caroline",
        total_score=7.5,
        xlsx_eval=xlsx_eval,
        reflection_eval=reflection_eval,
        flags_for_review=["Currency formatting confidence was low."],
    )


class FeedbackSheetWriterTests(unittest.TestCase):
    def test_feedback_template_key(self) -> None:
        self.assertEqual(
            feedback_template_key("Feedback_Sheet_Bagley_Caroline.docx"),
            "bagleycaroline",
        )

    def test_write_feedback_sheet_for_student_populates_table(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            templates_dir = root / "Feedback Sheets"
            templates_dir.mkdir(parents=True, exist_ok=True)

            template_path = templates_dir / "Feedback_Sheet_Bagley_Caroline.docx"
            output_path = root / "Filled" / "Feedback_Sheet_Bagley_Caroline.docx"
            build_template(template_path)

            used_template = write_feedback_sheet_for_student(
                student_key="bagleycaroline",
                grade=sample_grade(),
                feedback_sheets_dir=templates_dir,
                output_path=output_path,
            )

            self.assertEqual(used_template, template_path)
            self.assertTrue(output_path.exists())

            doc = Document(output_path)
            table = doc.tables[0]

            self.assertEqual(table.rows[1].cells[2].text, "1.0")
            self.assertEqual(table.rows[1].cells[3].text, "")

            self.assertEqual(table.rows[2].cells[2].text, "3.5")
            self.assertIn("Remove duplicate transaction rows", table.rows[2].cells[3].text)
            self.assertNotIn("Review flags", table.rows[2].cells[3].text)

            self.assertEqual(table.rows[3].cells[3].text, "")
            self.assertEqual(table.rows[4].cells[3].text, "")

            self.assertEqual(table.rows[5].cells[2].text, "1.25")
            self.assertIn("Meets Expectations", table.rows[5].cells[3].text)

            self.assertEqual(table.rows[6].cells[2].text, "0.75")
            self.assertIn("Source pages: 2-3", table.rows[6].cells[3].text)

            self.assertEqual(table.rows[7].cells[2].text, "7.5")
            self.assertEqual(table.rows[7].cells[3].text, "")

    def test_write_feedback_sheets_for_batch_tracks_missing_templates(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            templates_dir = root / "Feedback Sheets"
            output_dir = root / "Filled"
            templates_dir.mkdir(parents=True, exist_ok=True)

            build_template(templates_dir / "Feedback_Sheet_Bagley_Caroline.docx")

            grades = {
                "bagleycaroline": sample_grade(),
                "missingstudent": sample_grade(),
            }

            written, missing = write_feedback_sheets_for_batch(
                grades=grades,
                feedback_sheets_dir=templates_dir,
                output_dir=output_dir,
            )

            self.assertEqual(written, 1)
            self.assertEqual(missing, ["missingstudent"])
            self.assertTrue((output_dir / "Feedback_Sheet_Bagley_Caroline.docx").exists())

    def test_to_student_facing_step1b_notes_maps_and_falls_back(self) -> None:
        notes = to_student_facing_step1b_notes(
            [
                "Required 17-column structure is incorrect (missing required columns). Missing: OpenDate",
                "Duplicate transaction rows remain in the cleaned dataset. Found 20 duplicate row(s).",
                "Unrecognized rule text.",
                "No errors found.",
            ]
        )

        self.assertIn("Include these columns in your cleaned dataset: OpenDate.", notes)
        self.assertIn("Remove duplicate transaction rows so each transaction appears only once.", notes)
        self.assertIn("Review this section and make the needed updates in your cleaned dataset.", notes)
        self.assertIn("No updates needed for this section.", notes)

    def test_full_credit_rows_have_blank_feedback_notes(self) -> None:
        xlsx_eval = XlsxEvaluation(
            score_1a=1.0,
            notes_1a="All seven worksheets present.",
            score_1b=4.0,
            errors_1b=["No errors found."],
            review_flags_1b=["Some review flag that should not appear."],
            score_1c=0.5,
            notes_1c="Pivot object exists.",
            score_1d=0.5,
            notes_1d="Pivot object exists.",
            review_flags=[],
        )
        reflection_eval = ReflectionEvaluation(
            q1=ReflectionQuestionResult(
                score=2.0,
                label="Excellent",
                rationale="Complete and specific.",
            ),
            q2=ReflectionQuestionResult(
                score=2.0,
                label="Excellent",
                rationale="Complete and specific.",
            ),
            q2_start_page=1,
            q2_end_page=1,
            review_flags=[],
        )
        grade = FinalGrade(
            student_name="Perfect Student",
            total_score=10.0,
            xlsx_eval=xlsx_eval,
            reflection_eval=reflection_eval,
            flags_for_review=["Should not appear in feedback sheet."],
        )

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            templates_dir = root / "Feedback Sheets"
            templates_dir.mkdir(parents=True, exist_ok=True)

            template_path = templates_dir / "Feedback_Sheet_Perfect_Student.docx"
            output_path = root / "Filled" / "Feedback_Sheet_Perfect_Student.docx"
            build_template(template_path)

            used_template = write_feedback_sheet_for_student(
                student_key="perfectstudent",
                grade=grade,
                feedback_sheets_dir=templates_dir,
                output_path=output_path,
            )

            self.assertEqual(used_template, template_path)
            doc = Document(output_path)
            table = doc.tables[0]

            self.assertEqual(table.rows[1].cells[3].text, "")
            self.assertEqual(table.rows[2].cells[3].text, "")
            self.assertEqual(table.rows[3].cells[3].text, "")
            self.assertEqual(table.rows[4].cells[3].text, "")
            self.assertEqual(table.rows[5].cells[3].text, "")
            self.assertEqual(table.rows[6].cells[3].text, "")
            self.assertEqual(table.rows[7].cells[3].text, "")


if __name__ == "__main__":
    unittest.main()